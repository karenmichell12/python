from docx import Document
import pandas as pd
import unicodedata

file = input("Ingrese la ruta del archivo: ")

document = Document(file)

document.paragraphs  # to extract paragraphs
document.paragraphs[2].text  # gives the text

ruta1 = input("Ingrese la ruta donde se encuentra el excel: ")
    
df = pd.read_excel(ruta1) 
Lista = pd.Series.tolist(df["PALABRAS"])

#palabra = input("Ingrese la palabra a buscar: ")

def normalize(c):
    return unicodedata.normalize("NFD",c)[0]


contador = {}
for par in document.paragraphs:  # to extract the whole text
    for i in range(0,len(Lista)):
        R = (''.join(Lista[i]))
        if R in par.text.upper():
            print(par.text)
            R = ''.join(normalize(c) for c in str(R))
            R = R.encode("utf8").decode("ascii","ignore")
            if R in contador:
                contador[R] += 1
            else:
                contador[R] = 1
print(contador)
     


# # I tried the below code to find some specific term
# for i in range(0, 50, 1):
#   if (document.paragraphs[i].text == ('Some-word')):
#     print (document.paragraph)
    
